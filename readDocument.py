import time
import docx
import socket
import os

filename = "speakingTest.docx"


def send_text_over_udp(text, target_ip, target_port):
    # Create a UDP socket
    udp_socket = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)

    try:
        # Encode the text to bytes
        data = text.encode('utf-8')

        # Send the data over UDP
        udp_socket.sendto(data, (target_ip, target_port))
        print(f"Text sent successfully to {target_ip}:{target_port}")
    except socket.error as e:
        print(f"Error occurred while sending data: {e}")
    finally:
        # Close the socket
        udp_socket.close()


def delete_content_from_word_doc(file_path):
    os.remove(file_path)
    # Create a new empty Word document
    doc = docx.Document("clear.docx")

    # Save the new empty document, effectively replacing the original one
    doc.save(file_path)


# importing module
delay = 2
i = 1
prevlength = 0
Totaltime = 0
while (i > 0):
    doc = docx.Document(filename)
    all_paras = doc.paragraphs
    para = all_paras[0]
    # end = time.time()

    print(para.text)
    print(Totaltime)
    print("-------")

    length = len(para.text)
    time.sleep(delay)
    target_ip = "192.168.5.146"  # Replace this with the target IP address
    target_port = 5000  # Replace this with the target port number
    textSent = "text1:"+ para.text
    if length == 0:
        textSent = "text1:No Info"
    print(textSent)
    send_text_over_udp(textSent, target_ip, target_port)
# ctrl +delete on word document
